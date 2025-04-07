# -*- coding:utf-8 -*-

import os
from http import HTTPStatus
# 建议dashscope SDK 的版本 >= 1.20.11
from dashscope import Application
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import time


def main(f_input, f_question):
    print(u'第一阶段: 程序参数确认中......')
    time_current = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())
    print(u'当前系统时间:' + time_current)
    f_prompt, f_mode = query_general(f_input, f_question)
    if f_mode == u'文本检查':
        print(u'当前模式:' + f_mode)
        f_path, f_name = os.path.split(f_input)
        print(u'当前检查文件:' + f_name)
    elif f_mode == u'知识库问询':
        print(u'当前模式:' + f_mode)
        print(u'当前问题:' + f_question)
    else:
        print(u'未检测到要检查的文件或需要回答的问题，请重新确认参数')
        exit()
    print('--' * 6)
    result_txt = kw_query(f_prompt)
    result_output = kw_query_output(time_current, f_mode, f_input, f_question, result_txt)
    print(u'所有阶段均已完成，程序正常退出。退出代码' + str(result_output))


def query_general(f_input, f_question):
    # 把文件成果检查和简单提问都写成完成提示词prompt，以此为问题调用知识库进行检查
    if f_input:
        cont = read_docx(f_input)
        f_prompt = u'以最严格的标准，检查接下来冒号后的文字是否与知识库信息存在差异，并将差异之处逐一列举出来。重点检查词语和文字是否完全相同，对整段内容未完全匹配或内容上的缺省可以不列举出来:\n' + cont
        f_mode = u'文本检查'
    elif f_question:
        f_prompt = f_question + u'\n所有引用的文件必须写出含书名号的具体文件名称'
        f_mode = u'知识库问询'
    else:
        print(u'未检测到要检查的文件或需要回答的问题，无法正常生成提示词，请重新确认参数')
        exit()
    return f_prompt, f_mode


def kw_query(f_prompt):
    print(u'第二阶段: 联系知识库处理中......')
    time_period_start = time.time()
    # print(u'-----------文本修改模式专用调试语句，测试docx文件是否读取正确----------\n' + f_prompt)
    # return u'测试用返回值'
    response = Application.call(
        api_key="sk-",
        app_id='',
        prompt=f_prompt,
        rag_options={
            "pipeline_ids": ["fag3304t70"]
        }
        # 知识库可加载多个，以逗号分隔
    )

    if response.status_code != HTTPStatus.OK:
        print(f'request_id={response.request_id}')
        print(f'code={response.status_code}')
        print(f'message={response.message}')
        print(f'请参考文档：https://help.aliyun.com/zh/model-studio/developer-reference/error-code')
        exit()
    else:
        pass
        # print('%s\n' % (response.output.text))  # 处理只输出文本text
        # print('%s\n' % (response.usage))
    print(u'知识库检索阶段已完成.')
    print(u'第二阶段总计用时:' + str(int((time.time() - time_period_start) * 100) / 100) + u'秒.')
    print('--' * 6)
    return response.output.text


def kw_query_output(time_current, f_mode, f_input, f_question, result_txt):
    print(u'第三阶段: ' + f_mode + '结果输出......')
    time_period_start = time.time()
    with open('log.txt', 'a', encoding='utf-8') as f_log:
        f_log.write('\n'.join([
            time_current,
            '查询对象:' + f_input + f_question,
            result_txt
        ]).replace('\n\n', '\n') + '\n' * 2)
    print(f_mode + u'结果已写入同文件夹下的log.txt文件中.')

    if f_mode == u'文本检查':
        print(u'按错误对文本进行标黄......')
        print(u'EEEEEEEEEEEEEEEERROR：标黄功能还在想办法做，还没能实现')
        # f_doc = Document(f_input)
        # # 要查找并设置背景颜色的文本
        # target_text = "特定文字片段"
        #
        # # 遍历文档中的每个段落
        # for paragraph in f_doc.paragraphs:
        #     if target_text in paragraph.text:
        #         set_background_color(paragraph, target_text)
        #
        # # 保存修改后的文档
        # f_path, f_name = os.path.split(f_input)
        # f_doc.save(f_path + '\\' + f_name.split('.') + u'-订正.docx')
    print(u'第三阶段总计用时:' + str(int((time.time() - time_period_start) * 100) / 100) + u'秒.')
    print('--' * 6)
    return 0


def read_docx(file_path):
    doc = Document(file_path)
    content = []
    for para in doc.paragraphs:
        content.append(para.text)
    return '\n'.join(content)


def set_background_color(paragraph, text):
    for run in paragraph.runs:
        if text in run.text:
            # 创建一个新的 run 来替换原来的 run
            new_run = paragraph.add_run(run.text)
            new_run.style = run.style
            new_run.font.name = run.font.name
            new_run._r = run._r  # 复制原有的 XML 元素

            # 添加背景颜色
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '%02X%02X%02X' % (255, 255, 0))  # 设置背景颜色
            new_run._r.get_or_add_rPr().append(shading_elm)

            # 删除旧的 run
            paragraph._p.remove(run._r)


if __name__ == '__main__':
    # # 测试1：知识库问询
    # f_input = ''
    # f_question = u'习近平主席认为新型城镇化中存在哪些主要问题？'

    # # 测试3：知识库问询
    # f_input = ''
    # f_question = u'市级总规有哪些主要编制内容？'

    # 测试3：文本修改
    f_input = 'D:\\demos\\test.docx'
    f_question = ''
    main(f_input, f_question)
